#!/usr/bin/env python3
"""Convert docs/*.md to docs/word/*.docx for distribution."""

from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[1] / "docs"
OUT = ROOT / "word"

FILES = [
    ("USER_GUIDE.zh-TW.md", "USER_GUIDE.zh-TW.docx"),
    ("USER_GUIDE.en.md", "USER_GUIDE.en.docx"),
    ("TECHNICAL.zh-TW.md", "TECHNICAL.zh-TW.docx"),
    ("TECHNICAL.en.md", "TECHNICAL.en.docx"),
    ("README.md", "Documentation-Index.docx"),
]


def set_run_font(run, east_asia="Microsoft JhengHei", ascii_font="Calibri", size=11):
    run.font.name = ascii_font
    run.font.size = Pt(size)
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.get_or_add_rFonts()
    r_fonts.set(qn("w:eastAsia"), east_asia)


def add_formatted_runs(paragraph, text, base_size=11, bold=False):
    pattern = re.compile(r"(\*\*[^*]+\*\*|`[^`]+`|\[[^\]]+\]\([^)]+\))")
    pos = 0
    for match in pattern.finditer(text):
        if match.start() > pos:
            run = paragraph.add_run(text[pos : match.start()])
            set_run_font(run, size=base_size)
            if bold:
                run.bold = True
        token = match.group(0)
        if token.startswith("**") and token.endswith("**"):
            run = paragraph.add_run(token[2:-2])
            set_run_font(run, size=base_size)
            run.bold = True
        elif token.startswith("`") and token.endswith("`"):
            run = paragraph.add_run(token[1:-1])
            set_run_font(run, ascii_font="Consolas", east_asia="Consolas", size=base_size)
            run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
        else:
            label, url = re.match(r"\[([^\]]+)\]\(([^)]+)\)", token).groups()
            run = paragraph.add_run(f"{label} ({url})")
            set_run_font(run, size=base_size)
            run.font.color.rgb = RGBColor(0x05, 0x63, 0xC1)
        pos = match.end()
    if pos < len(text):
        run = paragraph.add_run(text[pos:])
        set_run_font(run, size=base_size)
        if bold:
            run.bold = True


def add_code_block(doc, code):
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(6)
    paragraph.paragraph_format.space_after = Pt(6)
    paragraph.paragraph_format.left_indent = Inches(0.2)
    run = paragraph.add_run(code.rstrip("\n"))
    set_run_font(run, ascii_font="Consolas", east_asia="Consolas", size=9)
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), "F5F5F5")
    shd.set(qn("w:val"), "clear")
    paragraph.paragraph_format.element.get_or_add_pPr().append(shd)


def add_table(doc, rows):
    if not rows:
        return
    cols = max(len(row) for row in rows)
    table = doc.add_table(rows=len(rows), cols=cols)
    table.style = "Table Grid"
    for i, row in enumerate(rows):
        for j in range(cols):
            cell = table.cell(i, j)
            cell.text = ""
            paragraph = cell.paragraphs[0]
            value = row[j] if j < len(row) else ""
            value = re.sub(r"\*\*([^*]+)\*\*", r"\1", value)
            value = re.sub(r"`([^`]+)`", r"\1", value)
            value = re.sub(r"\[([^\]]+)\]\([^)]+\)", r"\1", value)
            run = paragraph.add_run(value.strip())
            set_run_font(run, size=10)
            if i == 0:
                run.bold = True
    doc.add_paragraph()


def parse_table_block(lines, start):
    rows = []
    index = start
    while index < len(lines) and lines[index].strip().startswith("|"):
        line = lines[index].strip()
        if re.match(r"^\|[\s\-:|]+\|$", line):
            index += 1
            continue
        cells = [cell.strip() for cell in line.strip("|").split("|")]
        rows.append(cells)
        index += 1
    return rows, index


def md_to_docx(md_path: Path, docx_path: Path) -> None:
    lines = md_path.read_text(encoding="utf-8").splitlines()
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(0.9)
    section.right_margin = Inches(0.9)

    index = 0
    in_code = False
    code_lines: list[str] = []
    while index < len(lines):
        line = lines[index]
        if line.strip().startswith("```"):
            if not in_code:
                in_code = True
                code_lines = []
            else:
                add_code_block(doc, "\n".join(code_lines))
                in_code = False
            index += 1
            continue
        if in_code:
            code_lines.append(line)
            index += 1
            continue

        if (
            line.strip().startswith("|")
            and index + 1 < len(lines)
            and re.search(r"\|\s*-+", lines[index + 1])
        ):
            rows, index = parse_table_block(lines, index)
            add_table(doc, rows)
            continue

        if not line.strip():
            index += 1
            continue

        if line.startswith("# "):
            paragraph = doc.add_heading("", level=1)
            add_formatted_runs(paragraph, line[2:].strip(), base_size=18, bold=True)
        elif line.startswith("## "):
            paragraph = doc.add_heading("", level=2)
            add_formatted_runs(paragraph, line[3:].strip(), base_size=14, bold=True)
        elif line.startswith("### "):
            paragraph = doc.add_heading("", level=3)
            add_formatted_runs(paragraph, line[4:].strip(), base_size=12, bold=True)
        elif line.startswith("#### "):
            paragraph = doc.add_heading("", level=4)
            add_formatted_runs(paragraph, line[5:].strip(), base_size=11, bold=True)
        elif re.match(r"^[-*] ", line):
            paragraph = doc.add_paragraph(style="List Bullet")
            add_formatted_runs(paragraph, line[2:].strip())
        elif re.match(r"^\d+\. ", line):
            paragraph = doc.add_paragraph(style="List Number")
            add_formatted_runs(paragraph, re.sub(r"^\d+\. ", "", line).strip())
        elif line.strip() == "---":
            paragraph = doc.add_paragraph()
            p_pr = paragraph._p.get_or_add_pPr()
            p_bdr = OxmlElement("w:pBdr")
            bottom = OxmlElement("w:bottom")
            bottom.set(qn("w:val"), "single")
            bottom.set(qn("w:sz"), "6")
            bottom.set(qn("w:space"), "1")
            bottom.set(qn("w:color"), "AAAAAA")
            p_bdr.append(bottom)
            p_pr.append(p_bdr)
        elif line.startswith("> "):
            paragraph = doc.add_paragraph()
            paragraph.paragraph_format.left_indent = Inches(0.25)
            add_formatted_runs(paragraph, line[2:].strip())
            if paragraph.runs:
                paragraph.runs[0].italic = True
        else:
            paragraph = doc.add_paragraph()
            add_formatted_runs(paragraph, line.strip())
        index += 1

    docx_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(docx_path)


def main() -> int:
    OUT.mkdir(parents=True, exist_ok=True)
    for src_name, out_name in FILES:
        src = ROOT / src_name
        dst = OUT / out_name
        md_to_docx(src, dst)
        print(f"wrote {dst} ({dst.stat().st_size} bytes)")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
